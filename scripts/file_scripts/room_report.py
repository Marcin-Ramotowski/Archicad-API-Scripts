from collections import Counter
from data_tools.connection_init import acc, act, acu
from data_tools.functions import is_all_properties_available
from docx import Document
import ifcopenshell
from uuid import UUID


precision = 1
needed_properties = ['Zone_ZoneName', 'Zone_NetArea', ['Dane pomieszczenia', 'Kolor ścian'],
                    ['Specyfikacja produktu', 'Nazwa']]
all_properties = acc.GetAllPropertyNames()
run = is_all_properties_available(all_properties, needed_properties)

if run:
    ifc_file = ifcopenshell.open("others\ifc_files\model_testowy.ifc")

    def unpack_ifc_property(ifc_dict):
        values = []
        for ifc_type, ifc_keys in ifc_dict.items():
            parameters = ifc_file.by_type(ifc_type)[0]
            params_dict = parameters.get_info()
            values.append([params_dict[key] for key in ifc_keys])
        return values


    ifc_objects_to_get = {
        'IFCPROJECT' : ['Name', 'Phase'],
        'IFCSITE': ['SiteAddress'],
        'IFCPERSON': ['PrefixTitles', 'GivenName', 'FamilyName'],
        'IFCPROPERTYSET': ['HasProperties'] # stworzona przez autora grupa właściwości przydzielona do projektu
    }

    ifc_values = unpack_ifc_property(ifc_objects_to_get)

    # rozpakowywanie wartości
    project_params, address_params, person_params, custom_params = ifc_values
    project_name, phase = project_params
    address = ', '.join(address_params[0].get_info()['AddressLines'])
    if person_params[0] is None:
        designer = ' '.join(person_params[1:])
    else:
        person_params[0] = ' '.join(person_params[0])
        designer = ' '.join(person_params)

    unwrapped_customs = [param.get_info()['NominalValue'].get_info()['wrappedValue'] for param in custom_params[0]]
    project_number, start_date, end_date = unwrapped_customs

    # przygotowywanie listy pomieszczeń wraz z przydzielonymi do nich obiektami
    zone_properties = needed_properties[:3]
    object_properties = [needed_properties[3]]
    zone_properties_ids = [acu.GetBuiltInPropertyId(name) if type(name) is str else acu.GetUserDefinedPropertyId(*name)
                      for name in zone_properties]
    object_properties_ids = [acu.GetBuiltInPropertyId(name) if type(name) is str else acu.GetUserDefinedPropertyId(*name)
                      for name in object_properties]
    zones_ids = acc.GetElementsByType('Zone')
    zones = acc.GetElementsRelatedToZones(zones_ids, ['Object'])
    zones_descriptions = {}

    for i, zone in enumerate(zones):
        elements = zone.elements
        elements_values = [property.propertyValues[0].propertyValue.value for property in 
                           acc.GetPropertyValuesOfElements(elements, object_properties_ids)]
        elements_counter = Counter(elements_values).items()
        elements_values = [element if times == 1 else f"{element} x {times}" 
                           for element, times in elements_counter]
        cells_set = acc.GetPropertyValuesOfElements([zones_ids[i]], zone_properties_ids)
        zones_values = [act.ElementId(UUID(cell.value))for row in cells_set for cell in row][0]
        zones_values[1] = round(zones_values[1], precision)
        raport_values = zones_values[1:] + [elements_values]
        zones_descriptions[zones_values[0]] = raport_values

    # przygotowanie informacji do postaci gotowej do wyświetlenia w raporcie
    projects_data = [
    f"Raport z projektu nr {project_number}",
    f"Nazwa projektu: {project_name}",
    f"Faza projektu: {phase}",
    f"Lokalizacja: {address}",
    f"Data rozpoczęcia projektu: {start_date}",
    f"Data planowanego zakończenia projektu: {end_date}",
    f"Projektant: {designer}"]
    
    ''' Tworzenie raportu '''
    doc = Document()
    doc.add_heading(projects_data[0], level=1)

    # Dodawanie informacji ogólnych o projekcie
    doc.add_heading('1. Informacje ogólne o projekcie:', level=2)
    for row in projects_data[1:]:
        doc.add_paragraph(row)

    # Dodawanie informacji o pomieszczeniach
    doc.add_heading('2. Informacje o pomieszczeniach:', level=2)
    for zone, values in zones_descriptions.items():
        doc.add_paragraph(f"{zone}", style='List Number')
        doc.add_paragraph(f"Powierzchnia: {values[0]} m²")
        doc.add_paragraph(f"Kolor ścian: {values[1]}")
        doc.add_paragraph(f"Wyposażenie:")
        for object in values[2]:
            doc.add_paragraph(object, style='List Bullet')
        
    # Zapisywanie dokumentu do pliku .docx
    doc.save('raport_archicad.docx')
    acu.OpenFile('raport_archicad.docx')
